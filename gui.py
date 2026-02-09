import sys
import torch
import numpy as np
from config import *
from PIL import Image
from predict import Detector
import matplotlib.pyplot as plt
from scipy.interpolate import interp1d
from matplotlib.backends.backend_qt5agg import FigureCanvasQTAgg as FigureCanvas
from scipy.signal import find_peaks
from scipy.signal import savgol_filter
from PyQt5.QtCore import Qt
from PyQt5.QtGui import QPixmap, QTransform, QPen, QPainter, QImage, QCursor
from PyQt5.QtWidgets import (QMainWindow, QApplication, QWidget, QHBoxLayout, QSplitter, QVBoxLayout, QLabel,
                             QPushButton, QScrollArea, QTableWidget, QFileDialog, QMessageBox, QMenu, QAction,
                             QHeaderView, QTableWidgetItem, QSizePolicy)
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


class PeakAnalyzerApp(QMainWindow):
    def __init__(self):
        super().__init__()
        self.original_pixmap = None  # Main image in pixmap format
        self.main_image = None  # Main image in image format
        self.swimlane_split_line = None  # Swimlane dividing lines for generating swimlanes between pairs
        self.swimlane_location = None  # Position information of swimlanes (including start and end points)
        self.line_threshold = 10  # Threshold for mouse proximity to horizontal lines (pixels)
        self.is_dragging = False  # Global drag switch: whether dragging lines
        self.selected_line = -1  # Index of the dragged swimlane line
        self.selected_lane = -1  # Index of currently selected swimlane (-1 means unselected)
        self.selected_region = None
        self.swimlane_projection = None
        self.swimlane_pixmap = None
        self.swimlane_peak_projection = None
        self.swimlane_peak_pixmap = None
        self.peak_information = None
        self.selected_peak = None
        self.swimlane_peak_select_pixmap = None
        self.is_add_adjust_lane_mode = False  # Flag for "add swimlane dividing line mode" activation (False=not pressed, True=kept pressed)
        self.is_finish_lane_mode = False  # Flag for completion of swimlane adjustment
        self.is_adjust_region_mode = False  # Flag for target region selection
        self.is_finish_region_mode = False  # Flag for completion of target region selection
        self.detector = Detector()

        # Set window title and size
        self.setWindowTitle('Electrophoretic band detection')
        self.setGeometry(100, 100, 1200, 800)

        # Create central widget and main layout
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        main_layout = QHBoxLayout(central_widget)

        # Create left-right splitter
        splitter = QSplitter(Qt.Horizontal)
        main_layout.addWidget(splitter)

        # Left panel - Main image display area
        left_panel = QWidget()
        left_layout = QVBoxLayout(left_panel)

        # Main image display area
        self.main_image_container = QWidget()
        container_layout = QVBoxLayout(self.main_image_container)
        self.main_image_label = QLabel('Please load the main image')
        self.main_image_label.setAlignment(Qt.AlignCenter)
        self.main_image_label.setMinimumSize(400, 400)
        self.main_image_label.setStyleSheet("border: 1px solid #cccccc;")
        container_layout.addWidget(self.main_image_label)
        self.main_image_label.setMouseTracking(True)  # Critical! Enable mouseMoveEvent without pressing mouse
        self.main_image_label.mousePressEvent = self.on_image_click  # Bind mouse click event
        self.main_image_label.mouseMoveEvent = self.on_image_mouse_move  # Bind mouse drag event

        # Create button area with three rows
        buttons_container = QWidget()
        buttons_layout = QVBoxLayout(buttons_container)

        # First row button: Load main image
        row1_layout = QHBoxLayout()
        self.load_button = QPushButton('Load the main image')
        self.load_button.clicked[bool].connect(self.load_main_image)
        row1_layout.addWidget(self.load_button)
        buttons_layout.addLayout(row1_layout)

        # Second row buttons: Automatic swimlane segmentation, add swimlane dividing line, delete swimlane, select target region
        row2_layout = QHBoxLayout()
        self.split_lane_button = QPushButton('Lane segmentation')
        self.split_lane_button.clicked[bool].connect(self.perform_swimlane_split)
        self.finish_lane_button = QPushButton('End lane segmentation')
        self.finish_lane_button.clicked[bool].connect(self.finish_swimlane)
        self.select_region_button = QPushButton('Target area selection')
        self.select_region_button.clicked[bool].connect(self.perform_select_region)
        self.finish_region_button = QPushButton('End target area selection')
        self.finish_region_button.clicked[bool].connect(self.finish_select_region)

        # Disable these buttons initially
        self.split_lane_button.setEnabled(False)
        self.finish_lane_button.setEnabled(False)
        self.select_region_button.setEnabled(False)
        self.finish_region_button.setEnabled(False)

        row2_layout.addWidget(self.split_lane_button)
        row2_layout.addWidget(self.finish_lane_button)
        row2_layout.addWidget(self.select_region_button)
        row2_layout.addWidget(self.finish_region_button)
        buttons_layout.addLayout(row2_layout)

        # Third row button: Peak analysis
        row3_layout = QHBoxLayout()
        self.analyze_button = QPushButton('Peak analysis')
        self.analyze_button.clicked[bool].connect(self.perform_peak_analysis)
        self.analyze_button.setEnabled(False)  # Disable initially
        row3_layout.addWidget(self.analyze_button)
        buttons_layout.addLayout(row3_layout)

        # Create scroll area for scrolling when image size exceeds display area
        scroll_area = QScrollArea()
        scroll_area.setWidgetResizable(True)
        scroll_area.setWidget(self.main_image_container)

        # Right panel - Group image display, table and page control
        right_panel = QWidget()
        right_layout = QVBoxLayout(right_panel)

        self.group_label_label = QLabel(f'Lane {self.selected_lane + 1} Profile Display:')

        # Group image display area (two images arranged vertically)
        group_display = QWidget()
        group_layout = QVBoxLayout(group_display)
        group_layout.setSpacing(5)

        self.group_labels = [QLabel() for _ in range(2)]
        for i, label in enumerate(self.group_labels):
            label.setAlignment(Qt.AlignCenter)
            label.setStyleSheet("border: 1px solid #cccccc;")
            # Critical: Make widget expandable
            label.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Expanding)
            group_layout.addWidget(label)

        # Page control area
        page_control = QWidget()
        page_layout = QHBoxLayout(page_control)
        self.prev_button = QPushButton('Previous Lane')
        self.prev_button.clicked[bool].connect(self.prev_page)
        self.prev_button.setEnabled(False)  # Disable initially
        self.page_label = QLabel(f'Lane {self.selected_lane + 1}/{0}')
        self.page_label.setAlignment(Qt.AlignCenter)
        self.next_button = QPushButton('Next Lane')
        self.next_button.clicked[bool].connect(self.next_page)
        self.next_button.setEnabled(False)  # Disable initially
        page_layout.addWidget(self.prev_button)
        page_layout.addWidget(self.page_label)
        page_layout.addWidget(self.next_button)

        # Table area
        self.table_label_label = QLabel(f'Lane {self.selected_lane + 1} Data Table:')
        self.table_widget = QTableWidget()
        self.table_widget.setColumnCount(5)
        self.table_widget.setHorizontalHeaderLabels(['Position', 'Area', 'Peak Height', 'Left Width', 'Right Width'])
        header = self.table_widget.horizontalHeader()
        for column in range(5):
            header.setSectionResizeMode(column, QHeaderView.Stretch)
        # Critical: Make table expandable
        self.table_widget.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Expanding)
        self.table_widget.setMinimumHeight(150)

        # Generate report button
        self.report_button = QPushButton('Generate Report')
        self.report_button.clicked[bool].connect(self.generate_report)
        self.report_button.setEnabled(False)  # Disable initially

        left_layout.addWidget(scroll_area)
        left_layout.addWidget(buttons_container)

        right_layout.addWidget(self.group_label_label)
        right_layout.addWidget(group_display)
        right_layout.addWidget(page_control)
        right_layout.addWidget(self.table_label_label)
        right_layout.addWidget(self.table_widget)
        right_layout.addWidget(self.report_button)

        right_layout.setStretchFactor(self.group_label_label, 0)
        right_layout.setStretchFactor(group_display, 9)
        right_layout.setStretchFactor(page_control, 1)
        right_layout.setStretchFactor(self.table_label_label, 0)
        right_layout.setStretchFactor(self.table_widget, 3)
        right_layout.setStretchFactor(self.report_button, 1)

        splitter.addWidget(left_panel)
        splitter.addWidget(right_panel)

        splitter.setStretchFactor(0, 2)
        splitter.setStretchFactor(1, 4)

    def load_main_image(self):
        """Load main image and display on interface"""
        # Open file selection dialog to allow selecting image files
        file_path, _ = QFileDialog.getOpenFileName(
            None, "Select Image File", "",
            "Image Files (*.png *.jpg *.jpeg *.bmp *.tif *.tiff);;All Files (*)"
        )

        # Return if user cancels selection
        if not file_path:
            return

        try:
            # Load image and rotate 270 degrees
            self.original_pixmap = QPixmap(file_path)
            self.original_pixmap = self.original_pixmap.transformed(
                QTransform().rotate(270),
                Qt.SmoothTransformation
            )

            # Check if image loaded successfully
            if self.original_pixmap.isNull():
                QMessageBox.warning(None, "Load Failed", "Cannot load selected image, please check file format and permissions.")
                return

            # Save original image data for subsequent processing
            self.main_image = Image.open(file_path).convert('L').rotate(90, expand=True)

            # Display image on label
            self.main_image_label.setPixmap(self.original_pixmap.scaled(
                self.main_image_label.size(),
                Qt.KeepAspectRatio,
                Qt.SmoothTransformation
            ))

            # Enable swimlane segmentation button (image loaded), disable target region selection and peak analysis buttons
            self.split_lane_button.setEnabled(True)
            self.finish_lane_button.setEnabled(False)
            self.select_region_button.setEnabled(False)
            self.finish_region_button.setEnabled(False)
            self.analyze_button.setEnabled(False)
            self.prev_button.setEnabled(False)
            self.next_button.setEnabled(False)

            self.swimlane_split_line = None  # Swimlane dividing lines for generating swimlanes between pairs
            self.swimlane_location = None  # Position information of swimlanes (including start and end points)
            self.is_dragging = False  # Global drag switch: whether dragging lines
            self.selected_line = -1  # Index of the dragged swimlane line
            self.selected_lane = -1  # Index of currently selected swimlane (-1 means unselected)
            self.selected_region = None
            self.swimlane_projection = None
            self.swimlane_pixmap = None
            self.swimlane_peak_projection = None
            self.swimlane_peak_pixmap = None
            self.peak_information = None
            self.selected_peak = None
            self.swimlane_peak_select_pixmap = None
            self.is_add_adjust_lane_mode = False  # Flag for "add swimlane dividing line mode" activation (False=not pressed, True=kept pressed)
            self.is_finish_lane_mode = False  # Flag for completion of swimlane adjustment
            self.is_adjust_region_mode = False  # Flag for target region selection
            self.is_finish_region_mode = False  # Flag for completion of target region selection

        except Exception as e:
            QMessageBox.critical(self, "Error", f"Error loading image: {str(e)}")
            self.statusBar().showMessage("Image load failed")

    def perform_swimlane_split(self):
        def transfer_image_to_horizontal_projection(img):
            width, height = img.size
            row_sums = []
            width_max = width * 255
            for x in range(height):
                col_sum = 0
                for y in range(150):
                    col_sum += img.getpixel((y, x))
                row_sums.append(col_sum / width_max)

            return np.array(row_sums), width_max

        def get_gate_centers(curve_up, prominence=0.0001):
            """Get center point sequence of each gate from curve"""
            curve_np = savgol_filter(curve_up, 100, 3)
            peaks, _ = find_peaks(curve_np.tolist(), prominence=prominence)
            peaks = peaks.tolist()
            self.swimlane_split_line = peaks
            self.swimlane_location = list(zip(self.swimlane_split_line[:-1], self.swimlane_split_line[1:]))
            return curve_np

        projection, _ = transfer_image_to_horizontal_projection(self.main_image)
        _ = get_gate_centers(projection)
        # Display dividing lines and group images
        self.show_select_region_and_swimlane_split()
        # Enable target region selection button
        self.select_region_button.setEnabled(False)
        self.finish_lane_button.setEnabled(True)
        self.is_add_adjust_lane_mode = True

    def finish_swimlane(self):
        self.is_finish_lane_mode = True
        self.is_add_adjust_lane_mode = False
        self.split_lane_button.setStyleSheet("")
        self.split_lane_button.setEnabled(False)
        self.select_region_button.setEnabled(True)

    def perform_select_region(self):
        if not self.swimlane_location:
            QMessageBox.information(self, "Information", "Please perform swimlane segmentation first")
            return

        def transfer_image_to_vertical_projection(img):
            width, height = img.size
            col_sums = []
            for x in range(width):
                col_sum = 0
                for y in range(height):
                    col_sum += img.getpixel((x, y))
                col_sums.append(col_sum)
            col_sums = np.max(col_sums) + np.min(col_sums) - col_sums
            col_sums = savgol_filter(col_sums, 15, 3)
            return col_sums

        def detect_dark_band(sequence):
            peaks, properties = find_peaks(
                sequence,
                prominence=0.05,  # Peak prominence threshold
                width=2,  # Minimum peak width
                rel_height=0.5  # Relative height for calculating peak width (half height)
            )
            max_peak_idx = np.argmax(np.array(sequence)[peaks])
            dark_band = peaks[max_peak_idx] - (peaks[max_peak_idx] - properties['left_ips'][max_peak_idx]) * 3 / 1.17
            return dark_band

        sequence = transfer_image_to_vertical_projection(self.main_image)
        dark_band = detect_dark_band(sequence.tolist())
        pixmap = QPixmap(self.original_pixmap)
        self.selected_region = (200, int(dark_band)) if dark_band > 200 else (200, int(pixmap.width() - 100))
        self.show_select_region_and_swimlane_split()
        self.is_adjust_region_mode = True

        self.finish_region_button.setEnabled(True)
        self.analyze_button.setEnabled(False)
        self.prev_button.setEnabled(False)
        self.next_button.setEnabled(False)

    def finish_select_region(self):
        self.is_finish_region_mode = True
        self.is_adjust_region_mode = False
        self.select_region_button.setStyleSheet("")
        self.select_region_button.setEnabled(False)
        self.finish_region_button.setEnabled(True)
        self.analyze_button.setEnabled(True)

    def on_image_mouse_move(self, event):
        """Triggered when mouse moves over image (without clicking)"""
        # 1. Check if image and swimlane dividing lines exist (skip if either missing)
        if not self.original_pixmap or self.original_pixmap.isNull() or not (
                self.is_add_adjust_lane_mode or self.is_adjust_region_mode):
            # Restore default cursor if no image/lines
            self.main_image_label.setCursor(QCursor(Qt.ArrowCursor))
            return

        if self.is_add_adjust_lane_mode:
            # 2. Coordinate conversion: Convert label coordinates to original image coordinates (critical! Avoid position judgment errors due to scaling)
            # Calculate scale factor (original image height / label display height)
            scale_factor = self.original_pixmap.height() / self.main_image_label.height()
            # Label coordinates → Original image coordinates (only Y coordinate needed as swimlane lines are horizontal)
            label_y = event.pos().y()
            original_y = int(label_y * scale_factor)

            # 3. Initialize "line hit" flag
            hit_line = False
            # Iterate through all swimlane dividing lines to check mouse proximity
            for i, line_y in enumerate(self.swimlane_split_line):
                # Note: Compare original image line Y coordinate with converted mouse Y coordinate (ensure accurate positioning)
                if abs(line_y - original_y) <= self.line_threshold:
                    hit_line = True
                    self.selected_line = i  # Record selected line index (for potential dragging)
                    # Set up-down arrow cursor only on main image label (precise scope)
                    self.main_image_label.setCursor(QCursor(Qt.SizeVerCursor))
                    break  # Stop iteration after hitting one line

            # 4. Restore default cursor if no line is hit
            if not hit_line:
                self.selected_line = -1  # Reset selected line index
                self.main_image_label.setCursor(QCursor(Qt.ArrowCursor))

            # (Optional) Keep is_dragging logic for future line dragging implementation
            self.is_dragging = hit_line  # Mark as draggable only when line is hit
            if self.is_dragging and self.selected_line != -1 and (event.buttons() & Qt.LeftButton):
                # Coordinate conversion: Label coordinates -> Original image coordinates
                scale_factor = self.original_pixmap.height() / self.main_image_label.height()
                label_y = event.pos().y()
                original_y = int(label_y * scale_factor)
                # Update line position and redraw
                self.swimlane_split_line[self.selected_line] = original_y
                self.swimlane_split_line = sorted(list(set(self.swimlane_split_line)))
                self.swimlane_location = list(zip(self.swimlane_split_line[:-1], self.swimlane_split_line[1:]))
                self.show_select_region_and_swimlane_split()
        elif self.is_adjust_region_mode:
            # 2. Coordinate conversion: Convert label coordinates to original image coordinates (critical! Avoid position judgment errors due to scaling)
            # Calculate scale factor (original image width / label display width)
            scale_factor = self.original_pixmap.height() / self.main_image_label.height()
            x_offset = (self.main_image_label.width() - self.original_pixmap.width() / scale_factor) / 2
            # Label coordinates → Original image coordinates (only X coordinate needed)
            label_x = event.pos().x()
            original_x = int((label_x - x_offset) * scale_factor)
            # Display debug information in status bar
            # status_text = f"Label X: {label_x}, Scale Factor: {x_offset:.4f}, Original X: {original_x}, Image Width: {self.original_pixmap.height()}, Label Width: {self.main_image_label.width()}"
            # self.statusBar().showMessage(status_text)  # Show status bar information (persists by default)
            if abs(self.selected_region[0] - original_x) <= self.line_threshold:
                hit_line = True
                self.main_image_label.setCursor(QCursor(Qt.SizeHorCursor))  # Horizontal double arrow
                select_line = 0
            elif abs(self.selected_region[1] - original_x) <= self.line_threshold:
                hit_line = True
                self.main_image_label.setCursor(QCursor(Qt.SizeHorCursor))  # Horizontal double arrow
                select_line = 1
            else:
                hit_line = False
                select_line = -1
                self.main_image_label.setCursor(QCursor(Qt.ArrowCursor))

            self.is_dragging = hit_line  # Mark as draggable only when line is hit
            if self.is_dragging and select_line != -1 and (event.buttons() & Qt.LeftButton):
                scale_factor = self.original_pixmap.height() / self.main_image_label.height()
                x_offset = (self.main_image_label.width() - self.original_pixmap.width() / scale_factor) / 2
                # Label coordinates → Original image coordinates (only X coordinate needed)
                label_x = event.pos().x()
                original_x = int((label_x - x_offset) * scale_factor)
                self.selected_region = (original_x, self.selected_region[1]) if select_line == 0 else (
                self.selected_region[0], original_x)
                self.show_select_region_and_swimlane_split()

    def on_image_click(self, event):
        """Handle mouse click events on image to implement swimlane selection"""
        # Get Y coordinate of click position (relative to image)
        click_pos = event.pos()
        # Calculate scale factor (actual image height / display label height)
        scale_factor = self.original_pixmap.height() / self.main_image_label.height()
        # Convert to original image coordinates
        original_y = int(click_pos.y() * scale_factor)
        # print(original_y)
        # print(self.swimlane_location)
        if self.is_add_adjust_lane_mode and not self.is_dragging:
            if not hasattr(self, 'swimlane_split_line'):
                self.swimlane_split_line = []  # Initialize list (if not initialized)
            too_close = False
            for existing_y in self.swimlane_split_line:
                if abs(original_y - existing_y) <= self.line_threshold:
                    too_close = True
            if not too_close:
                # Distance meets requirements, add new line
                self.swimlane_split_line.append(original_y)
                # Remove duplicates and keep sorted (optional, sorting facilitates subsequent processing)
                self.swimlane_split_line = sorted(list(set(self.swimlane_split_line)))
            else:
                # Optional: Add prompt indicating line is too close
                print(f"Cannot add line, distance to existing line is less than {self.line_threshold} pixels")
            # Sort (ensure dividing lines are ordered by Y coordinate)
            self.swimlane_split_line.sort()
            self.swimlane_location = list(zip(self.swimlane_split_line[:-1], self.swimlane_split_line[1:]))
            self.show_select_region_and_swimlane_split()

        elif event.button() == Qt.LeftButton and self.swimlane_location and not self.is_add_adjust_lane_mode and not self.is_adjust_region_mode:
            # Find swimlane containing click position
            for idx, (start, end) in enumerate(self.swimlane_location):
                if start <= original_y <= end:
                    self.selected_lane = idx
                    self.show_select_region_and_swimlane_split()
                    if self.swimlane_pixmap:
                        self.update_group_display()
                    return

            # When no swimlane is hit
            self.selected_lane = -1

    def contextMenuEvent(self, event):
        """Handle right-click menu events"""
        # Show right-click menu only when swimlane is selected
        if self.selected_lane != -1 and self.main_image_label.geometry().contains(event.pos()):
            # Create right-click menu
            menu = QMenu(self)

            # Add delete swimlane option
            delete_action = QAction("Delete Selected Lane", self)
            delete_action.triggered[bool].connect(self.delete_selected_swimlane)
            menu.addAction(delete_action)

            # Show menu
            menu.exec_(event.globalPos())
        else:
            # Call base class method to handle other right-clicks
            super().contextMenuEvent(event)

    def keyPressEvent(self, event):
        """Handle keyboard key press events"""
        # When Delete key is pressed and swimlane is selected
        if event.key() == Qt.Key_Delete and self.selected_lane != -1:
            self.delete_selected_swimlane()
        else:
            # Call base class method to handle other keys
            super().keyPressEvent(event)

    def delete_selected_swimlane(self):
        if self.selected_lane == -1 or not self.swimlane_location or not self.is_finish_lane_mode:
            return
        else:
            self.split_lane_button.setEnabled(False)
            del self.swimlane_location[self.selected_lane]
            if self.swimlane_pixmap:
                del self.swimlane_pixmap[self.selected_lane]
            if self.swimlane_projection:
                del self.swimlane_projection[self.selected_lane]
            if self.swimlane_peak_projection:
                del self.swimlane_peak_projection[self.selected_lane]
            if self.swimlane_peak_pixmap:
                del self.swimlane_peak_pixmap[self.selected_lane]
            if self.peak_information:
                del self.peak_information[self.selected_lane]
            self.selected_lane = self.selected_lane if self.selected_lane < len(self.swimlane_location) else 0
            # print(self.selected_lane)
            self.show_select_region_and_swimlane_split()
            self.update_group_display()

    def update_group_display(self):
        if not self.swimlane_pixmap or not self.swimlane_peak_pixmap:
            return
        self.group_labels[0].setPixmap(self.swimlane_pixmap[self.selected_lane].scaled(
            self.group_labels[1].size(),
            Qt.KeepAspectRatio,
            Qt.SmoothTransformation))
        self.group_labels[1].setPixmap(self.swimlane_peak_pixmap[self.selected_lane].scaled(
            self.group_labels[1].size(),
            Qt.KeepAspectRatio,
            Qt.SmoothTransformation))
        self.group_label_label.setText(f"Lane {self.selected_lane + 1} Profile Display:")
        self.table_label_label.setText(f"Lane {self.selected_lane + 1} Data Table:")
        self.page_label.setText(f'Page {self.selected_lane + 1}/{len(self.swimlane_location)}')
        self.update_table()

    def prev_page(self):
        self.selected_lane = (self.selected_lane - 1) % len(self.swimlane_location)
        self.update_group_display()
        self.show_select_region_and_swimlane_split()

    def next_page(self):
        self.selected_lane = (self.selected_lane + 1) % len(self.swimlane_location)
        self.update_group_display()
        self.show_select_region_and_swimlane_split()

    def show_select_region_and_swimlane_split(self):
        if self.original_pixmap.isNull():
            return

        pixmap = QPixmap(self.original_pixmap)

        if self.selected_region:
            with QPainter(pixmap) as painter:
                # Set green pen with 2-pixel width
                blue_pen = QPen(Qt.blue, 2)
                painter.setPen(blue_pen)
                painter.drawLine(self.selected_region[0], 0, self.selected_region[0], pixmap.height())
                painter.drawLine(self.selected_region[1], 0, self.selected_region[1], pixmap.height())
        if self.swimlane_location:
            with QPainter(pixmap) as painter:
                # Set green pen with 2-pixel width
                green_pen = QPen(Qt.green, 2)
                red_pen = QPen(Qt.red, 3)

                # Iterate through each y coordinate in peaks to draw horizontal lines and swimlane information (except last line)
                for i, (top, bottom) in enumerate(self.swimlane_location):
                    if i != self.selected_lane:
                        # Draw horizontal lines (using green pen)
                        painter.setPen(green_pen)
                        painter.drawLine(0, top, pixmap.width(), top)
                        painter.drawLine(0, bottom, pixmap.width(), bottom)
                        # Draw swimlane information (green text)
                        painter.setPen(QPen(Qt.green))
                        painter.drawText(10, top + 25, f"Lane {i + 1}")  # y[0]+25: offset below line

                if self.selected_lane != -1 and self.selected_lane < len(self.swimlane_location):
                    painter.setPen(red_pen)
                    top, bottom = self.swimlane_location[self.selected_lane]
                    painter.drawLine(0, top, pixmap.width(), top)
                    painter.drawLine(0, bottom, pixmap.width(), bottom)
                    painter.drawText(10, top + 25, f"Lane {self.selected_lane + 1} (selected)")

        # Display image
        self.main_image_label.setPixmap(pixmap.scaled(
            self.main_image_label.size(),
            Qt.KeepAspectRatio,
            Qt.SmoothTransformation
        ))

    def perform_peak_analysis(self):
        if not self.swimlane_location or not self.selected_region:
            return
        self.prev_button.setEnabled(True)
        self.next_button.setEnabled(True)
        self.swimlane_projection = []
        for i in range(len(self.swimlane_location)):
            swimlane_projection = []
            for x in range(*self.selected_region):
                col_sum = 0
                for y in range(*self.swimlane_location[i]):
                    col_sum += self.main_image.getpixel((x, y))
                col_sum = col_sum / abs(self.swimlane_location[i][1] - self.swimlane_location[i][0]) / 255
                swimlane_projection.append(col_sum)
            self.swimlane_projection.append(swimlane_projection)

        def adjust_to_length(sequence, max_value=0.7):
            # Original x coordinates
            x_original = np.arange(len(sequence))
            # Calculate x-axis scaling factor (scale sequence to target length)
            k_x = DATA_LENGTH / len(sequence)
            # Target x coordinates (1024 points)
            x_new = np.linspace(0, len(sequence) - 1, DATA_LENGTH)
            # Use linear interpolation to adjust sequence to fixed length (1024)
            f_interp = interp1d(x_original, sequence, kind='linear', fill_value="extrapolate")
            new_data = f_interp(x_new)
            # Normalization (set maximum value to specified value)
            sequence = (new_data / np.max(new_data)) * max_value
            k_y = max_value / np.max(new_data)  # Calculate y-axis scaling factor
            # Convert to PyTorch tensor and adjust shape
            sequence = torch.tensor(sequence).float()
            return sequence, k_y, k_x

        def adjust_back_to_length(sequence, projection):
            # Original x coordinates
            x_original = np.arange(len(sequence))
            # Target x coordinates
            x_new = np.linspace(0, len(sequence) - 1, len(projection))
            # Use linear interpolation to adjust sequence to fixed length
            f_interp = interp1d(x_original, sequence, kind='linear', fill_value="extrapolate")
            new_data = f_interp(x_new)
            # Convert to PyTorch tensor and adjust shape
            sequence = torch.tensor(new_data).float()
            return sequence

        def get_output(sequence, k_y, k_x, width_max):
            sequence_clone = sequence.reshape(1, 1, -1)
            output = self.detector(sequence_clone, target=torch.zeros(1, 16, 4), k_y=k_y, k_x=k_x, width_max=width_max)
            output_sequence, locations, areas, amplitudes, left_widths, right_widths = output

            return output_sequence, locations, areas, amplitudes, left_widths, right_widths

        self.swimlane_peak_projection = []
        self.peak_information = []

        for i in range(len(self.swimlane_projection)):
            sequence, k_y, k_x = adjust_to_length(self.swimlane_projection[i])
            width_max = abs(self.swimlane_location[i][1] - self.swimlane_location[i][0]) * 255
            output, locations, areas, amplitudes, left_widths, right_widths = get_output(sequence, k_y, k_x,
                                                                                         width_max=width_max)
            peak_information = [
                [locations[j], areas[j], amplitudes[j], left_widths[j], right_widths[j]]
                for j in range(len(locations))
            ]
            outputs = []
            for j in range(len(output)):
                output_result = adjust_back_to_length(output[j], self.swimlane_projection[i])
                outputs.append(output_result)
            self.swimlane_peak_projection.append(outputs)
            self.peak_information.append(peak_information)

        self.selected_lane = 0
        self.show_select_region_and_swimlane_split()
        self.plot_curve(width=self.group_labels[0].width(), height=self.group_labels[0].height())
        self.plot_peak_curve(width=self.group_labels[1].width(), height=self.group_labels[1].height())
        self.update_group_display()
        self.report_button.setEnabled(True)

    def plot_curve(self, width=600, height=200):
        if not self.swimlane_projection:
            return
        self.swimlane_pixmap = []
        for i in range(len(self.swimlane_projection)):
            x = np.arange(len(self.swimlane_projection[i]))
            fig, ax = plt.subplots(figsize=(width / 100, height / 100), dpi=100)
            swimlane_projection = np.array(self.swimlane_projection[i]) * abs(
                self.swimlane_location[i][1] - self.swimlane_location[i][0]) * 255 / 1e4
            ax.plot(x, swimlane_projection, 'b-', linewidth=2)

            # Set plot properties
            ax.set_xlim(0, len(self.swimlane_projection[i]))  # Set x-axis range
            ax.set_ylim(0, 1.05 * max(swimlane_projection))  # Set y-axis range
            # Set x and y axis labels
            ax.set_xlabel('Location(Pixel)', fontsize=16, fontname='Times New Roman')  # x-axis label, font size 10px
            ax.set_ylabel('Grayscale(×10⁴)', fontsize=16, fontname='Times New Roman')  # y-axis label, font size 10px

            # Set x and y axis tick label sizes
            ax.tick_params(axis='x', labelsize=12)  # x-axis tick label size 8px
            ax.tick_params(axis='y', labelsize=12)  # y-axis tick label size 8px

            for label in ax.get_xticklabels():
                label.set_fontname('Times New Roman')
            for label in ax.get_yticklabels():
                label.set_fontname('Times New Roman')

            plt.tight_layout()  # Auto adjust layout

            # Draw to canvas
            canvas = FigureCanvas(fig)
            canvas.draw()
            # Convert to QPixmap
            width, height = canvas.get_width_height()
            image = QImage(canvas.buffer_rgba(), width, height, QImage.Format_ARGB32)
            pixmap = QPixmap.fromImage(image)

            plt.close(fig)

            self.swimlane_pixmap.append(pixmap)

    def plot_peak_curve(self, width=600, height=200):
        if not self.swimlane_peak_projection:
            return
        self.swimlane_peak_pixmap = []
        # Define a set of fill colors (extendable as needed)
        fill_colors = ['#FF6666', '#3399FF', '#66CC66', '#FFB366', '#FF66B2', '#B366FF']

        for i in range(len(self.swimlane_peak_projection)):
            fig, ax = plt.subplots(figsize=(width / 100, height / 100), dpi=100)
            swimlane_projection = np.array(self.swimlane_projection[i]) * abs(
                self.swimlane_location[i][1] - self.swimlane_location[i][0]) * 255

            for j in range(len(self.swimlane_peak_projection[i])):
                # Get curve data for current peak
                curve_data = self.swimlane_peak_projection[i][j] / 1e4
                # Generate x-axis coordinates (matching curve length)
                x = np.arange(len(curve_data))
                # Plot curve
                ax.plot(x, curve_data, linewidth=2,
                        color=fill_colors[j % len(fill_colors)],  # Curve color matches fill color
                        label=f'Peak {j + 1}')  # Optional: Add label for legend
                # Fill area under curve (between curve and x-axis)
                ax.fill_between(x, curve_data, 0,
                                color=fill_colors[j % len(fill_colors)],  # Use cycled colors
                                alpha=1.0)  # Transparency (0-1, smaller = more transparent)

            # Set plot properties
            ax.set_xlim(0, len(self.swimlane_projection[i]))
            ax.set_ylim(0, 1.05 * max(swimlane_projection) / 1e4)

            # Set x and y axis labels
            ax.set_xlabel('Location(Pixel)', fontsize=16, fontname='Times New Roman')
            ax.set_ylabel('Grayscale(×10⁴)', fontsize=16, fontname='Times New Roman')

            # Set tick labels
            ax.tick_params(axis='x', labelsize=12)
            ax.tick_params(axis='y', labelsize=12)
            for label in ax.get_xticklabels() + ax.get_yticklabels():
                label.set_fontname('Times New Roman')

            # Optional: Add legend (if peak identification is needed)
            # ax.legend(fontsize=16, prop={'family': 'Times New Roman'})

            plt.tight_layout()

            # Draw to canvas and convert to QPixmap (keep original logic)
            canvas = FigureCanvas(fig)
            canvas.draw()
            width, height = canvas.get_width_height()
            image = QImage(canvas.buffer_rgba(), width, height, QImage.Format_ARGB32)
            pixmap = QPixmap.fromImage(image)
            plt.close(fig)

            self.swimlane_peak_pixmap.append(pixmap)

    def update_table(self):
        # Update table data
        if self.selected_lane < len(self.peak_information):
            self.table_widget.setRowCount(len(self.peak_information[self.selected_lane]))
            for row, data in enumerate(self.peak_information[self.selected_lane]):
                for col, value in enumerate(data):
                    item = QTableWidgetItem(f"{value:.3f}")
                    item.setTextAlignment(Qt.AlignCenter)
                    self.table_widget.setItem(row, col, item)

    def generate_report(self):
        """Generate table data of all swimlanes into a Word document"""
        if not self.peak_information or len(self.peak_information) == 0:
            QMessageBox.information(self, "Information", "No data available for report generation, please complete peak analysis first")
            return

        # Open file save dialog to get save path
        file_path, _ = QFileDialog.getSaveFileName(
            self, "Save Report", "", "Word Documents (*.docx);;All Files (*)"
        )
        if not file_path:
            return  # User cancels save

        try:
            # Create new Word document
            doc = Document()

            # Add document title
            title = doc.add_heading('Electrophoretic Band Detection Report', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Center title

            # Add report description
            doc.add_paragraph('This report contains peak analysis data for each swimlane, including position, area, peak height, left width and right width.')
            doc.add_paragraph('')  # Empty line separator

            # Iterate through all swimlanes to generate tables
            for lane_idx in range(len(self.peak_information)):
                # Add swimlane title (level 2 heading)
                lane_heading = doc.add_heading(f'Swimlane {lane_idx + 1} Data Analysis', level=2)
                lane_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

                # Get peak information for current swimlane
                peak_data = self.peak_information[lane_idx]
                if not peak_data:
                    doc.add_paragraph('No detected peak data for this swimlane.')
                    doc.add_paragraph('')  # Empty line separator
                    continue

                # Create table (rows = data rows + 1 header row, columns = 5 columns)
                table = doc.add_table(rows=1, cols=5)
                table.autofit = False  # Disable auto-fit, use fixed width
                table.allow_autofit = False

                # Set table column widths (total width ~6 inches, adjustable as needed)
                col_widths = [1.2, 1.2, 1.2, 1.2, 1.2]  # Column widths (inches)
                for i in range(5):
                    table.columns[i].width = Inches(col_widths[i])

                # Fill table header row
                headers = ['Position', 'Area', 'Peak Height', 'Left Width', 'Right Width']
                hdr_cells = table.rows[0].cells
                for i, header in enumerate(headers):
                    hdr_cells[i].text = header
                    # Set header cell style (bold, centered)
                    self._set_cell_style(hdr_cells[i], bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

                # Fill table data rows
                for row_data in peak_data:
                    row_cells = table.add_row().cells
                    for i, value in enumerate(row_data):
                        # Display with 3 decimal places
                        row_cells[i].text = f"{value:.3f}"
                        # Set data cell style (centered)
                        self._set_cell_style(row_cells[i], align=WD_ALIGN_PARAGRAPH.CENTER)

                # Add empty line separator between swimlane tables
                doc.add_paragraph('')

            # Save document
            doc.save(file_path)
            QMessageBox.information(self, "Success", f"Report generated and saved successfully to:\n{file_path}")

        except Exception as e:
            QMessageBox.critical(self, "Error", f"Error generating report:\n{str(e)}")

    @staticmethod
    def _set_cell_style(cell, bold: bool = False, align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.LEFT):
        """Set table cell style (public interface only, no protected member access)"""
        paragraph = cell.paragraphs[0]
        paragraph.alignment = align

        # Ensure run object exists
        if not paragraph.runs:
            paragraph.add_run()
        run = paragraph.runs[0]

        # Critical: Set Western + Chinese font names first (using docx auto-association)
        run.font.name = "SimSun"  # Western identifier for SimSun
        run.font.size = Pt(10)
        run.bold = bold

        # Force trigger Chinese font rendering (no need for _element)
        # Add empty run with Chinese character to let docx auto-detect Chinese font
        run_cn = paragraph.add_run("　")  # Full-width space, no display impact
        run_cn.font.name = "SimSun"
        run_cn.font.size = Pt(10)
        run_cn.bold = bold

        # Remove temporarily added full-width space (keep only font configuration)
        run_cn.text = ""


def main():
    app = QApplication(sys.argv)
    window = PeakAnalyzerApp()
    window.show()
    window.showMaximized()  # Maximize window
    sys.exit(app.exec_())


if __name__ == '__main__':
    main()