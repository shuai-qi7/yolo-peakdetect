import sys
import torch
import numpy as np
from config import *
from PIL import Image
from predict import Detector
import matplotlib.pyplot as plt
from scipy.interpolate import interp1d
from matplotlib.backends.backend_qt5agg import FigureCanvasQTAgg as FigureCanvas
from scipy.signal import find_peaks
from scipy.signal import savgol_filter
from PyQt5.QtCore import Qt
from PyQt5.QtGui import QPixmap, QTransform, QPen, QPainter, QImage, QCursor
from PyQt5.QtWidgets import (QMainWindow, QApplication, QWidget, QHBoxLayout, QSplitter, QVBoxLayout, QLabel,
                             QPushButton, QScrollArea, QTableWidget, QFileDialog, QMessageBox, QMenu, QAction,
                             QHeaderView, QTableWidgetItem, QSizePolicy)
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


class PeakAnalyzerApp(QMainWindow):
    def __init__(self):
        super().__init__()
        self.original_pixmap = None  # pixmap类型的主图象
        self.main_image = None  # image类型的主图象
        self.swimlane_split_line = None  # 泳道分割线，用于两两之间生成泳道
        self.swimlane_location = None  # 泳道的位置信息（包括起点和终点）
        self.line_threshold = 10  # 鼠标靠近横线的阈值（像素）
        self.is_dragging = False  # 是否正在拖动线条（全局拖动开关）
        self.selected_line = -1  # 拖动泳道线索引
        self.selected_lane = -1  # 当前选中的泳道索引(-1表示未选中)
        self.selected_region = None
        self.swimlane_projection = None
        self.swimlane_pixmap = None
        self.swimlane_peak_projection = None
        self.swimlane_peak_pixmap = None
        self.peak_information = None
        self.selected_peak = None
        self.swimlane_peak_select_pixmap = None
        self.is_add_adjust_lane_mode = False  # 标记“增加泳道分割线模式”是否激活（False=未按下，True=按下保持）
        self.is_finish_lane_mode = False  # 标记是否完成泳道调整
        self.is_adjust_region_mode = False  # 标记是否进行目标区域选择
        self.is_finish_region_mode = False  # 标记是否完成目标区域选择
        self.detector = Detector()

        # 设置窗口标题和大小
        self.setWindowTitle('Electrophoretic band detection')
        self.setGeometry(100, 100, 1200, 800)

        # 创建中央部件和主布局
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        main_layout = QHBoxLayout(central_widget)

        # 创建左右分割器
        splitter = QSplitter(Qt.Horizontal)
        main_layout.addWidget(splitter)

        # 左侧面板 - 主图像显示区域
        left_panel = QWidget()
        left_layout = QVBoxLayout(left_panel)

        # 主图像显示区域
        self.main_image_container = QWidget()
        container_layout = QVBoxLayout(self.main_image_container)
        self.main_image_label = QLabel('Please load the main image')
        self.main_image_label.setAlignment(Qt.AlignCenter)
        self.main_image_label.setMinimumSize(400, 400)
        self.main_image_label.setStyleSheet("border: 1px solid #cccccc;")
        container_layout.addWidget(self.main_image_label)
        self.main_image_label.setMouseTracking(True)  # 关键！让标签不按鼠标也能触发mouseMoveEvent
        self.main_image_label.mousePressEvent = self.on_image_click  # 绑定鼠标点击事件
        self.main_image_label.mouseMoveEvent = self.on_image_mouse_move  # 绑定鼠标拖动事件

        # 创建按钮区域，分为三行
        buttons_container = QWidget()
        buttons_layout = QVBoxLayout(buttons_container)

        # 第一行按钮：加载主图像
        row1_layout = QHBoxLayout()
        self.load_button = QPushButton('Load the main image')
        self.load_button.clicked[bool].connect(self.load_main_image)
        row1_layout.addWidget(self.load_button)
        buttons_layout.addLayout(row1_layout)

        # 第二行按钮：自动泳道分割、增加泳道分割线、删除泳道、选择目标区域
        row2_layout = QHBoxLayout()
        self.split_lane_button = QPushButton('Lane segmentation')
        self.split_lane_button.clicked[bool].connect(self.perform_swimlane_split)
        self.finish_lane_button = QPushButton('End lane segmentation')
        self.finish_lane_button.clicked[bool].connect(self.finish_swimlane)
        self.select_region_button = QPushButton('Target area selection')
        self.select_region_button.clicked[bool].connect(self.perform_select_region)
        self.finish_region_button = QPushButton('End target area selection')
        self.finish_region_button.clicked[bool].connect(self.finish_select_region)

        # 初始禁用这些按钮
        self.split_lane_button.setEnabled(False)
        self.finish_lane_button.setEnabled(False)
        self.select_region_button.setEnabled(False)
        self.finish_region_button.setEnabled(False)

        row2_layout.addWidget(self.split_lane_button)
        row2_layout.addWidget(self.finish_lane_button)
        row2_layout.addWidget(self.select_region_button)
        row2_layout.addWidget(self.finish_region_button)
        buttons_layout.addLayout(row2_layout)

        # 第三行按钮：峰分析
        row3_layout = QHBoxLayout()
        self.analyze_button = QPushButton('Peak analysis')
        self.analyze_button.clicked[bool].connect(self.perform_peak_analysis)
        self.analyze_button.setEnabled(False)  # 初始禁用
        row3_layout.addWidget(self.analyze_button)
        buttons_layout.addLayout(row3_layout)

        # 创建一个滚动区域控件，用于在图像尺寸超过显示区域时提供滚动浏览功能
        scroll_area = QScrollArea()
        scroll_area.setWidgetResizable(True)
        scroll_area.setWidget(self.main_image_container)

        # 右侧面板 - 组图展示、表格和翻页控制
        right_panel = QWidget()
        right_layout = QVBoxLayout(right_panel)

        self.group_label_label = QLabel(f'Lane {self.selected_lane + 1} Profile Display:')

        # 组图展示区域（上下排列两张图片）
        group_display = QWidget()
        group_layout = QVBoxLayout(group_display)
        group_layout.setSpacing(5)

        self.group_labels = [QLabel() for _ in range(2)]
        for i, label in enumerate(self.group_labels):
            label.setAlignment(Qt.AlignCenter)
            label.setStyleSheet("border: 1px solid #cccccc;")
            # 关键：让部件可扩展
            label.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Expanding)
            group_layout.addWidget(label)

        # 翻页控制区域
        page_control = QWidget()
        page_layout = QHBoxLayout(page_control)
        self.prev_button = QPushButton('Previous Lane')
        self.prev_button.clicked[bool].connect(self.prev_page)
        self.prev_button.setEnabled(False)  # 初始禁用
        self.page_label = QLabel(f'Lane {self.selected_lane + 1}/{0}')
        self.page_label.setAlignment(Qt.AlignCenter)
        self.next_button = QPushButton('Next Lane')
        self.next_button.clicked[bool].connect(self.next_page)
        self.next_button.setEnabled(False)  # 初始禁用
        page_layout.addWidget(self.prev_button)
        page_layout.addWidget(self.page_label)
        page_layout.addWidget(self.next_button)

        # 表格区域
        self.table_label_label = QLabel(f'Lane {self.selected_lane + 1} Data Table:')
        self.table_widget = QTableWidget()
        self.table_widget.setColumnCount(5)
        self.table_widget.setHorizontalHeaderLabels(['Position', 'Area', 'Peak Height', 'Left Width', 'Right Width'])
        header = self.table_widget.horizontalHeader()
        for column in range(5):
            header.setSectionResizeMode(column, QHeaderView.Stretch)
        # 关键：让表格可扩展
        self.table_widget.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Expanding)
        self.table_widget.setMinimumHeight(150)

        # 生成报告按钮
        self.report_button = QPushButton('Generate Report')
        self.report_button.clicked[bool].connect(self.generate_report)
        self.report_button.setEnabled(False)  # 初始禁用

        left_layout.addWidget(scroll_area)
        left_layout.addWidget(buttons_container)

        right_layout.addWidget(self.group_label_label)
        right_layout.addWidget(group_display)
        right_layout.addWidget(page_control)
        right_layout.addWidget(self.table_label_label)
        right_layout.addWidget(self.table_widget)
        right_layout.addWidget(self.report_button)

        right_layout.setStretchFactor(self.group_label_label, 0)
        right_layout.setStretchFactor(group_display, 9)
        right_layout.setStretchFactor(page_control, 1)
        right_layout.setStretchFactor(self.table_label_label, 0)
        right_layout.setStretchFactor(self.table_widget, 3)
        right_layout.setStretchFactor(self.report_button, 1)

        splitter.addWidget(left_panel)
        splitter.addWidget(right_panel)

        splitter.setStretchFactor(0, 2)
        splitter.setStretchFactor(1, 4)

    def load_main_image(self):
        """加载主图像并显示在界面上"""
        # 打开文件选择对话框，允许选择图像文件
        file_path, _ = QFileDialog.getOpenFileName(
            None, "选择图像文件", "",
            "图像文件 (*.png *.jpg *.jpeg *.bmp *.tif *.tiff);;所有文件 (*)"
        )

        # 如果用户取消了选择，则返回
        if not file_path:
            return

        try:
            # 加载图像，并旋转270度
            self.original_pixmap = QPixmap(file_path)
            self.original_pixmap = self.original_pixmap.transformed(
                QTransform().rotate(270),
                Qt.SmoothTransformation
            )

            # 检查图像是否成功加载
            if self.original_pixmap.isNull():
                QMessageBox.warning(None, "加载失败", "无法加载所选图像，请检查文件格式和权限。")
                return

            # 保存原始图像数据，用于后续处理
            self.main_image = Image.open(file_path).convert('L').rotate(90, expand=True)

            # 在标签上显示图像
            self.main_image_label.setPixmap(self.original_pixmap.scaled(
                self.main_image_label.size(),
                Qt.KeepAspectRatio,
                Qt.SmoothTransformation
            ))

            # 启用泳道分割按钮，因为已经加载了图像，关闭选择目标区域按钮和峰分析按钮
            self.split_lane_button.setEnabled(True)
            self.finish_lane_button.setEnabled(False)
            self.select_region_button.setEnabled(False)
            self.finish_region_button.setEnabled(False)
            self.analyze_button.setEnabled(False)
            self.prev_button.setEnabled(False)
            self.next_button.setEnabled(False)

            self.swimlane_split_line = None  # 泳道分割线，用于两两之间生成泳道
            self.swimlane_location = None  # 泳道的位置信息（包括起点和终点）
            self.is_dragging = False  # 是否正在拖动线条（全局拖动开关）
            self.selected_line = -1  # 拖动泳道线索引
            self.selected_lane = -1  # 当前选中的泳道索引(-1表示未选中)
            self.selected_region = None
            self.swimlane_projection = None
            self.swimlane_pixmap = None
            self.swimlane_peak_projection = None
            self.swimlane_peak_pixmap = None
            self.peak_information = None
            self.selected_peak = None
            self.swimlane_peak_select_pixmap = None
            self.is_add_adjust_lane_mode = False  # 标记“增加泳道分割线模式”是否激活（False=未按下，True=按下保持）
            self.is_finish_lane_mode = False  # 标记是否完成泳道调整
            self.is_adjust_region_mode = False  # 标记是否进行目标区域选择
            self.is_finish_region_mode = False  # 标记是否完成目标区域选择

        except Exception as e:
            QMessageBox.critical(self, "错误", f"加载图像时发生错误: {str(e)}")
            self.statusBar().showMessage("加载图像失败")

    def perform_swimlane_split(self):
        def transfer_image_to_horizontal_projection(img):
            width, height = img.size
            row_sums = []
            width_max = width * 255
            for x in range(height):
                col_sum = 0
                for y in range(150):
                    col_sum += img.getpixel((y, x))
                row_sums.append(col_sum / width_max)

            return np.array(row_sums), width_max

        def get_gate_centers(curve_up, prominence=0.0001):
            """从曲线中获取每个门的中心点序列"""
            curve_np = savgol_filter(curve_up, 100, 3)
            peaks, _ = find_peaks(curve_np.tolist(), prominence=prominence)
            peaks = peaks.tolist()
            self.swimlane_split_line = peaks
            self.swimlane_location = list(zip(self.swimlane_split_line[:-1], self.swimlane_split_line[1:]))
            return curve_np

        projection, _ = transfer_image_to_horizontal_projection(self.main_image)
        _ = get_gate_centers(projection)
        # 显示分割线和组图
        self.show_select_region_and_swimlane_split()
        # 启用选择目标区域按钮
        self.select_region_button.setEnabled(False)
        self.finish_lane_button.setEnabled(True)
        self.is_add_adjust_lane_mode = True

    def finish_swimlane(self):
        self.is_finish_lane_mode = True
        self.is_add_adjust_lane_mode = False
        self.split_lane_button.setStyleSheet("")
        self.split_lane_button.setEnabled(False)
        self.select_region_button.setEnabled(True)

    def perform_select_region(self):
        if not self.swimlane_location:
            QMessageBox.information(self, "提示", "请先进行泳道分割")
            return

        def transfer_image_to_vertical_projection(img):
            width, height = img.size
            col_sums = []
            for x in range(width):
                col_sum = 0
                for y in range(height):
                    col_sum += img.getpixel((x, y))
                col_sums.append(col_sum)
            col_sums = np.max(col_sums) + np.min(col_sums) - col_sums
            col_sums = savgol_filter(col_sums, 15, 3)
            return col_sums

        def detect_dark_band(sequence):
            peaks, properties = find_peaks(
                sequence,
                prominence=0.05,  # 峰的突出度阈值
                width=2,  # 峰的最小宽度
                rel_height=0.5  # 计算峰宽的相对高度（半高宽）
            )
            max_peak_idx = np.argmax(np.array(sequence)[peaks])
            dark_band = peaks[max_peak_idx] - (peaks[max_peak_idx] - properties['left_ips'][max_peak_idx]) * 3 / 1.17
            return dark_band

        sequence = transfer_image_to_vertical_projection(self.main_image)
        dark_band = detect_dark_band(sequence.tolist())
        pixmap = QPixmap(self.original_pixmap)
        self.selected_region = (200, int(dark_band)) if dark_band > 200 else (200, int(pixmap.width() - 100))
        self.show_select_region_and_swimlane_split()
        self.is_adjust_region_mode = True

        self.finish_region_button.setEnabled(True)
        self.analyze_button.setEnabled(False)
        self.prev_button.setEnabled(False)
        self.next_button.setEnabled(False)

    def finish_select_region(self):
        self.is_finish_region_mode = True
        self.is_adjust_region_mode = False
        self.select_region_button.setStyleSheet("")
        self.select_region_button.setEnabled(False)
        self.finish_region_button.setEnabled(True)
        self.analyze_button.setEnabled(True)

    def on_image_mouse_move(self, event):
        """鼠标在图像上单纯移动（未点击）时触发"""
        # 1. 检查是否有图像和泳道分割线（缺一则不执行后续逻辑）
        if not self.original_pixmap or self.original_pixmap.isNull() or not (
                self.is_add_adjust_lane_mode or self.is_adjust_region_mode):
            # 若没有图像/线条，恢复主图像标签的默认光标
            self.main_image_label.setCursor(QCursor(Qt.ArrowCursor))
            return

        if self.is_add_adjust_lane_mode:
            # 2. 坐标转换：将标签坐标转为原始图像坐标（关键！避免缩放导致位置判断错误）
            # 计算缩放比例（原始图像高度 / 标签显示高度）
            scale_factor = self.original_pixmap.height() / self.main_image_label.height()
            # 标签坐标 → 原始图像坐标（仅需Y坐标，因为泳道线是水平的）
            label_y = event.pos().y()
            original_y = int(label_y * scale_factor)

            # 3. 初始化“是否命中线条”的标记
            hit_line = False
            # 遍历所有泳道分割线，判断鼠标是否靠近
            for i, line_y in enumerate(self.swimlane_split_line):
                # 注意：这里用原始图像的线条Y坐标和转换后的鼠标Y坐标比较（确保位置精准）
                if abs(line_y - original_y) <= self.line_threshold:
                    hit_line = True
                    self.selected_line = i  # 记录选中的线条索引（后续拖动可能用到）
                    # 仅在主图像标签上设置上下箭头光标（作用范围精准）
                    self.main_image_label.setCursor(QCursor(Qt.SizeVerCursor))
                    break  # 命中一条即可，无需继续遍历

            # 4. 若未命中任何线条，恢复主图像标签的默认光标
            if not hit_line:
                self.selected_line = -1  # 重置选中的线条索引
                self.main_image_label.setCursor(QCursor(Qt.ArrowCursor))

            # （可选）如果后续需要处理“拖动线条”，可保留is_dragging逻辑，此处暂不涉及
            self.is_dragging = hit_line  # 仅当命中线条时，标记为可拖动状态
            if self.is_dragging and self.selected_line != -1 and (event.buttons() & Qt.LeftButton):
                # 坐标转换：标签坐标 -> 原始图像坐标
                scale_factor = self.original_pixmap.height() / self.main_image_label.height()
                label_y = event.pos().y()
                original_y = int(label_y * scale_factor)
                # 更新线条位置并重新绘制
                self.swimlane_split_line[self.selected_line] = original_y
                self.swimlane_split_line = sorted(list(set(self.swimlane_split_line)))
                self.swimlane_location = list(zip(self.swimlane_split_line[:-1], self.swimlane_split_line[1:]))
                self.show_select_region_and_swimlane_split()
        elif self.is_adjust_region_mode:
            # 2. 坐标转换：将标签坐标转为原始图像坐标（关键！避免缩放导致位置判断错误）
            # 计算缩放比例（原始图像宽度 / 标签显示宽度）
            scale_factor = self.original_pixmap.height() / self.main_image_label.height()
            x_offset = (self.main_image_label.width() - self.original_pixmap.width() / scale_factor) / 2
            # 标签坐标 → 原始图像坐标（仅需X坐标）
            label_x = event.pos().x()
            original_x = int((label_x - x_offset) * scale_factor)
            # 在状态栏显示调试信息
            # status_text = f"标签X: {label_x}, 缩放比例: {x_offset:.4f}, 原始X: {original_x}，图片宽度：{self.original_pixmap.height()},标签宽度：{self.main_image_label.width()}"
            # self.statusBar().showMessage(status_text)  # 显示状态栏信息，默认会保持一段时间
            if abs(self.selected_region[0] - original_x) <= self.line_threshold:
                hit_line = True
                self.main_image_label.setCursor(QCursor(Qt.SizeHorCursor))  # 横向双箭头
                select_line = 0
            elif abs(self.selected_region[1] - original_x) <= self.line_threshold:
                hit_line = True
                self.main_image_label.setCursor(QCursor(Qt.SizeHorCursor))  # 横向双箭头
                select_line = 1
            else:
                hit_line = False
                select_line = -1
                self.main_image_label.setCursor(QCursor(Qt.ArrowCursor))

            self.is_dragging = hit_line  # 仅当命中线条时，标记为可拖动状态
            if self.is_dragging and select_line != -1 and (event.buttons() & Qt.LeftButton):
                scale_factor = self.original_pixmap.height() / self.main_image_label.height()
                x_offset = (self.main_image_label.width() - self.original_pixmap.width() / scale_factor) / 2
                # 标签坐标 → 原始图像坐标（仅需X坐标）
                label_x = event.pos().x()
                original_x = int((label_x - x_offset) * scale_factor)
                self.selected_region = (original_x, self.selected_region[1]) if select_line == 0 else (
                self.selected_region[0], original_x)
                self.show_select_region_and_swimlane_split()

    def on_image_click(self, event):
        """处理图像上的鼠标点击事件，实现泳道选择功能"""
        # 获取点击位置的Y坐标(相对于图像)
        click_pos = event.pos()
        # 计算缩放比例(实际图像高度 / 显示标签高度)
        scale_factor = self.original_pixmap.height() / self.main_image_label.height()
        # 转换为原始图像坐标
        original_y = int(click_pos.y() * scale_factor)
        # print(original_y)
        # print(self.swimlane_location)
        if self.is_add_adjust_lane_mode and not self.is_dragging:
            if not hasattr(self, 'swimlane_split_line'):
                self.swimlane_split_line = []  # 初始化列表（如果未初始化）
            too_close = False
            for existing_y in self.swimlane_split_line:
                if abs(original_y - existing_y) <= self.line_threshold:
                    too_close = True
            if not too_close:
                # 距离符合要求，添加新线条
                self.swimlane_split_line.append(original_y)
                # 去重并保持排序（可选，排序便于后续处理）
                self.swimlane_split_line = sorted(list(set(self.swimlane_split_line)))
            else:
                # 可以添加提示信息，说明线条过近
                print(f"无法添加线条，与已有线条距离小于{self.line_threshold}像素")
            # 排序（确保分割线按Y坐标从小到大排列）
            self.swimlane_split_line.sort()
            self.swimlane_location = list(zip(self.swimlane_split_line[:-1], self.swimlane_split_line[1:]))
            self.show_select_region_and_swimlane_split()

        elif event.button() == Qt.LeftButton and self.swimlane_location and not self.is_add_adjust_lane_mode and not self.is_adjust_region_mode:
            # 查找点击位置所在的泳道
            for idx, (start, end) in enumerate(self.swimlane_location):
                if start <= original_y <= end:
                    self.selected_lane = idx
                    self.show_select_region_and_swimlane_split()
                    if self.swimlane_pixmap:
                        self.update_group_display()
                    return

            # 未命中任何泳道时
            self.selected_lane = -1

    def contextMenuEvent(self, event):
        """处理右键菜单事件"""
        # 仅在已选中泳道时显示右键菜单
        if self.selected_lane != -1 and self.main_image_label.geometry().contains(event.pos()):
            # 创建右键菜单
            menu = QMenu(self)

            # 添加删除泳道选项
            delete_action = QAction("删除选中泳道", self)
            delete_action.triggered[bool].connect(self.delete_selected_swimlane)
            menu.addAction(delete_action)

            # 显示菜单
            menu.exec_(event.globalPos())
        else:
            # 调用基类方法处理其他右键点击
            super().contextMenuEvent(event)

    def keyPressEvent(self, event):
        """处理键盘按键事件"""
        # 当Delete键被按下且已选中泳道时
        if event.key() == Qt.Key_Delete and self.selected_lane != -1:
            self.delete_selected_swimlane()
        else:
            # 调用基类方法处理其他按键
            super().keyPressEvent(event)

    def delete_selected_swimlane(self):
        if self.selected_lane == -1 or not self.swimlane_location or not self.is_finish_lane_mode:
            return
        else:
            self.split_lane_button.setEnabled(False)
            del self.swimlane_location[self.selected_lane]
            if self.swimlane_pixmap:
                del self.swimlane_pixmap[self.selected_lane]
            if self.swimlane_projection:
                del self.swimlane_projection[self.selected_lane]
            if self.swimlane_peak_projection:
                del self.swimlane_peak_projection[self.selected_lane]
            if self.swimlane_peak_pixmap:
                del self.swimlane_peak_pixmap[self.selected_lane]
            if self.peak_information:
                del self.peak_information[self.selected_lane]
            self.selected_lane = self.selected_lane if self.selected_lane < len(self.swimlane_location) else 0
            # print(self.selected_lane)
            self.show_select_region_and_swimlane_split()
            self.update_group_display()

    def update_group_display(self):
        if not self.swimlane_pixmap or not self.swimlane_peak_pixmap:
            return
        self.group_labels[0].setPixmap(self.swimlane_pixmap[self.selected_lane].scaled(
            self.group_labels[1].size(),
            Qt.KeepAspectRatio,
            Qt.SmoothTransformation))
        self.group_labels[1].setPixmap(self.swimlane_peak_pixmap[self.selected_lane].scaled(
            self.group_labels[1].size(),
            Qt.KeepAspectRatio,
            Qt.SmoothTransformation))
        self.group_label_label.setText(f"Lane {self.selected_lane + 1} Profile Display:")
        self.table_label_label.setText(f"Lane {self.selected_lane + 1} Data Table:")
        self.page_label.setText(f'第 {self.selected_lane + 1}/{len(self.swimlane_location)} 页')
        self.update_table()

    def prev_page(self):
        self.selected_lane = (self.selected_lane - 1) % len(self.swimlane_location)
        self.update_group_display()
        self.show_select_region_and_swimlane_split()

    def next_page(self):
        self.selected_lane = (self.selected_lane + 1) % len(self.swimlane_location)
        self.update_group_display()
        self.show_select_region_and_swimlane_split()

    def show_select_region_and_swimlane_split(self):
        if self.original_pixmap.isNull():
            return

        pixmap = QPixmap(self.original_pixmap)

        if self.selected_region:
            with QPainter(pixmap) as painter:
                # 设置绿色画笔，线宽2像素
                blue_pen = QPen(Qt.blue, 2)
                painter.setPen(blue_pen)
                painter.drawLine(self.selected_region[0], 0, self.selected_region[0], pixmap.height())
                painter.drawLine(self.selected_region[1], 0, self.selected_region[1], pixmap.height())
        if self.swimlane_location:
            with QPainter(pixmap) as painter:
                # 设置绿色画笔，线宽2像素
                green_pen = QPen(Qt.green, 2)
                red_pen = QPen(Qt.red, 3)

                # 遍历peaks中的每个y坐标，绘制水平线和泳道信息（最后一条除外）
                for i, (top, bottom) in enumerate(self.swimlane_location):
                    if i != self.selected_lane:
                        # 绘制水平线（使用绿色画笔）
                        painter.setPen(green_pen)
                        painter.drawLine(0, top, pixmap.width(), top)
                        painter.drawLine(0, bottom, pixmap.width(), bottom)
                        # 绘制泳道信息（绿色文本）
                        painter.setPen(QPen(Qt.green))
                        painter.drawText(10, top + 25, f"Lane {i + 1}")  # y[0]+25：线下方偏移

                if self.selected_lane != -1 and self.selected_lane < len(self.swimlane_location):
                    painter.setPen(red_pen)
                    top, bottom = self.swimlane_location[self.selected_lane]
                    painter.drawLine(0, top, pixmap.width(), top)
                    painter.drawLine(0, bottom, pixmap.width(), bottom)
                    painter.drawText(10, top + 25, f"Lane {self.selected_lane + 1} (selected)")

        # 显示图片
        self.main_image_label.setPixmap(pixmap.scaled(
            self.main_image_label.size(),
            Qt.KeepAspectRatio,
            Qt.SmoothTransformation
        ))

    def perform_peak_analysis(self):
        if not self.swimlane_location or not self.selected_region:
            return
        self.prev_button.setEnabled(True)
        self.next_button.setEnabled(True)
        self.swimlane_projection = []
        for i in range(len(self.swimlane_location)):
            swimlane_projection = []
            for x in range(*self.selected_region):
                col_sum = 0
                for y in range(*self.swimlane_location[i]):
                    col_sum += self.main_image.getpixel((x, y))
                col_sum = col_sum / abs(self.swimlane_location[i][1] - self.swimlane_location[i][0]) / 255
                swimlane_projection.append(col_sum)
            self.swimlane_projection.append(swimlane_projection)

        def adjust_to_length(sequence, max_value=0.7):
            # 原始x坐标
            x_original = np.arange(len(sequence))
            # 计算x轴缩放因子（将序列缩放到长度）
            k_x = DATA_LENGTH / len(sequence)
            # 目标x坐标（1024个点）
            x_new = np.linspace(0, len(sequence) - 1, DATA_LENGTH)
            # 使用线性插值将序列调整为固定长度（1024）
            f_interp = interp1d(x_original, sequence, kind='linear', fill_value="extrapolate")
            new_data = f_interp(x_new)
            # 归一化处理（使最大值为指定值）
            sequence = (new_data / np.max(new_data)) * max_value
            k_y = max_value / np.max(new_data)  # 计算y轴缩放因子
            # 转换为PyTorch张量并调整形状
            sequence = torch.tensor(sequence).float()
            return sequence, k_y, k_x

        def adjust_back_to_length(sequence, projection):
            # 原始x坐标
            x_original = np.arange(len(sequence))
            # 目标x坐标
            x_new = np.linspace(0, len(sequence) - 1, len(projection))
            # 使用线性插值将序列调整为固定长度
            f_interp = interp1d(x_original, sequence, kind='linear', fill_value="extrapolate")
            new_data = f_interp(x_new)
            # 转换为PyTorch张量并调整形状
            sequence = torch.tensor(new_data).float()
            return sequence

        def get_output(sequence, k_y, k_x, width_max):
            sequence_clone = sequence.reshape(1, 1, -1)
            output = self.detector(sequence_clone, target=torch.zeros(1, 16, 4), k_y=k_y, k_x=k_x, width_max=width_max)
            output_sequence, locations, areas, amplitudes, left_widths, right_widths = output

            return output_sequence, locations, areas, amplitudes, left_widths, right_widths

        self.swimlane_peak_projection = []
        self.peak_information = []

        for i in range(len(self.swimlane_projection)):
            sequence, k_y, k_x = adjust_to_length(self.swimlane_projection[i])
            width_max = abs(self.swimlane_location[i][1] - self.swimlane_location[i][0]) * 255
            output, locations, areas, amplitudes, left_widths, right_widths = get_output(sequence, k_y, k_x,
                                                                                         width_max=width_max)
            peak_information = [
                [locations[j], areas[j], amplitudes[j], left_widths[j], right_widths[j]]
                for j in range(len(locations))
            ]
            outputs = []
            for j in range(len(output)):
                output_result = adjust_back_to_length(output[j], self.swimlane_projection[i])
                outputs.append(output_result)
            self.swimlane_peak_projection.append(outputs)
            self.peak_information.append(peak_information)

        self.selected_lane = 0
        self.show_select_region_and_swimlane_split()
        self.plot_curve(width=self.group_labels[0].width(), height=self.group_labels[0].height())
        self.plot_peak_curve(width=self.group_labels[1].width(), height=self.group_labels[1].height())
        self.update_group_display()
        self.report_button.setEnabled(True)

    def plot_curve(self, width=600, height=200):
        if not self.swimlane_projection:
            return
        self.swimlane_pixmap = []
        for i in range(len(self.swimlane_projection)):
            x = np.arange(len(self.swimlane_projection[i]))
            fig, ax = plt.subplots(figsize=(width / 100, height / 100), dpi=100)
            swimlane_projection = np.array(self.swimlane_projection[i]) * abs(
                self.swimlane_location[i][1] - self.swimlane_location[i][0]) * 255 / 1e4
            ax.plot(x, swimlane_projection, 'b-', linewidth=2)

            # 设置图形属性
            ax.set_xlim(0, len(self.swimlane_projection[i]))  # 设置x轴范围
            ax.set_ylim(0, 1.05 * max(swimlane_projection))  # 设置y轴范围
            # 设置x、y轴标题
            ax.set_xlabel('Location(Pixel)', fontsize=16, fontname='Times New Roman')  # x轴标题，字体大小10px
            ax.set_ylabel('Grayscale(×10⁴)', fontsize=16, fontname='Times New Roman')  # y轴标题，字体大小10px

            # 设置x、y轴刻度文字大小
            ax.tick_params(axis='x', labelsize=12)  # x轴刻度文字大小8px
            ax.tick_params(axis='y', labelsize=12)  # y轴刻度文字大小8px

            for label in ax.get_xticklabels():
                label.set_fontname('Times New Roman')
            for label in ax.get_yticklabels():
                label.set_fontname('Times New Roman')

            plt.tight_layout()  # 自动调整布局

            # 绘制到画布
            canvas = FigureCanvas(fig)
            canvas.draw()
            # 转换为QPixmap
            width, height = canvas.get_width_height()
            image = QImage(canvas.buffer_rgba(), width, height, QImage.Format_ARGB32)
            pixmap = QPixmap.fromImage(image)

            plt.close(fig)

            self.swimlane_pixmap.append(pixmap)

    def plot_peak_curve(self, width=600, height=200):
        if not self.swimlane_peak_projection:
            return
        self.swimlane_peak_pixmap = []
        # 定义一组填充颜色（可根据需要扩展）
        fill_colors = ['#FF6666', '#3399FF', '#66CC66', '#FFB366', '#FF66B2', '#B366FF']

        for i in range(len(self.swimlane_peak_projection)):
            fig, ax = plt.subplots(figsize=(width / 100, height / 100), dpi=100)
            swimlane_projection = np.array(self.swimlane_projection[i]) * abs(
                self.swimlane_location[i][1] - self.swimlane_location[i][0]) * 255

            for j in range(len(self.swimlane_peak_projection[i])):
                # 获取当前峰的曲线数据
                curve_data = self.swimlane_peak_projection[i][j] / 1e4
                # 生成x轴坐标（与曲线长度一致）
                x = np.arange(len(curve_data))
                # 绘制曲线
                ax.plot(x, curve_data, linewidth=2,
                        color=fill_colors[j % len(fill_colors)],  # 曲线颜色与填充色一致
                        label=f'Peak {j + 1}')  # 可选：添加标签用于图例
                # 填充曲线下方区域（曲线与x轴之间）
                ax.fill_between(x, curve_data, 0,
                                color=fill_colors[j % len(fill_colors)],  # 使用循环的颜色
                                alpha=1.0)  # 透明度（0-1之间，值越小越透明）

            # 设置图形属性
            ax.set_xlim(0, len(self.swimlane_projection[i]))
            ax.set_ylim(0, 1.05 * max(swimlane_projection) / 1e4)

            # 设置x、y轴标题
            ax.set_xlabel('Location(Pixel)', fontsize=16, fontname='Times New Roman')
            ax.set_ylabel('Grayscale(×10⁴)', fontsize=16, fontname='Times New Roman')

            # 设置刻度文字
            ax.tick_params(axis='x', labelsize=12)
            ax.tick_params(axis='y', labelsize=12)
            for label in ax.get_xticklabels() + ax.get_yticklabels():
                label.set_fontname('Times New Roman')

            # 可选：添加图例（如果需要显示每个峰的标识）
            # ax.legend(fontsize=16, prop={'family': 'Times New Roman'})

            plt.tight_layout()

            # 绘制到画布并转换为QPixmap（保持原有逻辑）
            canvas = FigureCanvas(fig)
            canvas.draw()
            width, height = canvas.get_width_height()
            image = QImage(canvas.buffer_rgba(), width, height, QImage.Format_ARGB32)
            pixmap = QPixmap.fromImage(image)
            plt.close(fig)

            self.swimlane_peak_pixmap.append(pixmap)

    def update_table(self):
        # 更新表格数据
        if self.selected_lane < len(self.peak_information):
            self.table_widget.setRowCount(len(self.peak_information[self.selected_lane]))
            for row, data in enumerate(self.peak_information[self.selected_lane]):
                for col, value in enumerate(data):
                    item = QTableWidgetItem(f"{value:.3f}")
                    item.setTextAlignment(Qt.AlignCenter)
                    self.table_widget.setItem(row, col, item)

    def generate_report(self):
        """将所有泳道的表格数据生成到一个Word文档中"""
        if not self.peak_information or len(self.peak_information) == 0:
            QMessageBox.information(self, "提示", "没有可生成报告的数据，请先完成峰分析")
            return

        # 打开文件保存对话框，获取保存路径
        file_path, _ = QFileDialog.getSaveFileName(
            self, "保存报告", "", "Word文档 (*.docx);;所有文件 (*)"
        )
        if not file_path:
            return  # 用户取消保存

        try:
            # 创建一个新的Word文档
            doc = Document()

            # 添加文档标题
            title = doc.add_heading('电泳条带检测报告', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 标题居中

            # 添加报告说明
            doc.add_paragraph('本报告包含各泳道的峰分析数据，具体包括位置、面积、峰高、左侧宽度和右侧宽度等信息。')
            doc.add_paragraph('')  # 空行分隔

            # 遍历所有泳道，生成表格
            for lane_idx in range(len(self.peak_information)):
                # 添加泳道标题（二级标题）
                lane_heading = doc.add_heading(f'泳道 {lane_idx + 1} 数据分析', level=2)
                lane_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

                # 获取当前泳道的峰信息
                peak_data = self.peak_information[lane_idx]
                if not peak_data:
                    doc.add_paragraph('该泳道无检测到的峰数据。')
                    doc.add_paragraph('')  # 空行分隔
                    continue

                # 创建表格（行数=数据行数+1行标题，列数=5列）
                table = doc.add_table(rows=1, cols=5)
                table.autofit = False  # 禁用自动调整，使用固定宽度
                table.allow_autofit = False

                # 设置表格列宽（总宽度约6英寸，可根据需要调整）
                col_widths = [1.2, 1.2, 1.2, 1.2, 1.2]  # 每列宽度（英寸）
                for i in range(5):
                    table.columns[i].width = Inches(col_widths[i])

                # 填充表格标题行
                headers = ['位置', '面积', '峰高', '左侧宽度', '右侧宽度']
                hdr_cells = table.rows[0].cells
                for i, header in enumerate(headers):
                    hdr_cells[i].text = header
                    # 设置标题单元格样式（加粗、居中）
                    self._set_cell_style(hdr_cells[i], bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

                # 填充表格数据行
                for row_data in peak_data:
                    row_cells = table.add_row().cells
                    for i, value in enumerate(row_data):
                        # 保留3位小数显示
                        row_cells[i].text = f"{value:.3f}"
                        # 设置数据单元格样式（居中）
                        self._set_cell_style(row_cells[i], align=WD_ALIGN_PARAGRAPH.CENTER)

                # 添加空行分隔不同泳道的表格
                doc.add_paragraph('')

            # 保存文档
            doc.save(file_path)
            QMessageBox.information(self, "成功", f"报告已成功生成并保存至:\n{file_path}")

        except Exception as e:
            QMessageBox.critical(self, "错误", f"生成报告时发生错误:\n{str(e)}")

    @staticmethod
    def _set_cell_style(cell, bold: bool = False, align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.LEFT):
        """设置表格单元格样式（纯公开接口，无受保护成员访问）"""
        paragraph = cell.paragraphs[0]
        paragraph.alignment = align

        # 确保有运行对象
        if not paragraph.runs:
            paragraph.add_run()
        run = paragraph.runs[0]

        # 核心：先统一设置西文+中文字体名称（利用 docx 自动关联）
        run.font.name = "SimSun"  # 宋体的西文标识
        run.font.size = Pt(10)
        run.bold = bold

        # 强制触发中文字体渲染（无需 _element）
        # 添加一个含中文字符的空 run，让 docx 自动识别中文字体
        run_cn = paragraph.add_run("　")  # 全角空格，不影响显示
        run_cn.font.name = "宋体"
        run_cn.font.size = Pt(10)
        run_cn.bold = bold

        # 移除临时添加的全角空格（只保留字体配置）
        run_cn.text = ""


def main():
    app = QApplication(sys.argv)
    window = PeakAnalyzerApp()
    window.show()
    window.showMaximized()  # 最大化窗口
    sys.exit(app.exec_())


if __name__ == '__main__':
    main()
